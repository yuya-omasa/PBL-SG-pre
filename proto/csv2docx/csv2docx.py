import docx
import csv

#Word文書の新規作成
doc = docx.Document()

doc.add_heading('台本')

#段落を追加する
doc.add_paragraph('2024/12/01')

#csvファイルを読み込み
with open('sample.csv') as f:
    reader = csv.reader(f)
    l = [row for row in reader]

doc.add_table(rows=len(l), cols=len(l[0]))
tbl = doc.tables[0]
r = 0
c = 0
for row in tbl.rows:
    for cell in row.cells:
        cell.add_paragraph(l[r][c])
        c += 1
    c = 0
    r += 1


#Word文書の保存
doc.save("Sample.docx")

#参考　https://gammasoft.jp/support/how-to-use-python-docx-for-word-file/